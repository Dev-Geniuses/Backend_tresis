import difflib
import copy
from docx import Document
from docx.shared import RGBColor

def resaltar_diferencias(doc1_path, doc2_path):
    doc_original = Document(doc1_path)
    doc_copia = copy.deepcopy(doc_original)

    doc2 = Document(doc2_path)

    if len(doc_original.paragraphs) >= len(doc2.paragraphs):
        doc_largo = doc_original
        doc_corto = doc2
    else:
        doc_largo = doc2
        doc_corto = doc_original

    diferencias = []

    for i in range(len(doc_largo.paragraphs)):
        paragraph_largo = doc_largo.paragraphs[i].text

        if i < len(doc_corto.paragraphs):
            paragraph_corto = doc_corto.paragraphs[i].text
        else:
            paragraph_corto = ""

        diff = difflib.ndiff(paragraph_largo.split(), paragraph_corto.split())

        for j, d in enumerate(diff):
            if d.startswith('- '):
                diferencias.append({
                    "parrafo": i,
                    "palabra": d[2:],
                    "posicion": j
                })

    for diferencia in diferencias:
        parrafo_index = diferencia["parrafo"]
        palabra_index = diferencia["posicion"]
        paragraph = doc_copia.paragraphs[parrafo_index]

        if palabra_index < len(paragraph.runs):
            run = paragraph.runs[palabra_index]
            font = run.font
            font.color.rgb = RGBColor(255, 0, 0)  # Resaltar en rojo la palabra

    doc_copia.save("./documents/documento1_comparado.docx")
    if diferencias:
        return f"Se han encontrado diferencias. Las diferencias se han guardado en el documento: {output_path}."
    else:
        return "No se han encontrado diferencias entre los documentos."

documento1_path = "documento1.docx"
documento2_path = "documento2.docx"
output_path = "documento1_resaltado.docx"

def resaltar_diferencias_documento2(doc1_path, doc2_path):
    doc1 = Document(doc1_path)
    doc2 = Document(doc2_path)
    doc_copia = doc2

    diff = difflib.ndiff(
        [p.text for p in doc2.paragraphs],
        [p.text for p in doc1.paragraphs]
    )

    diferencias = []

    for i, d in enumerate(diff):
        if d.startswith('- '):
            diferencias.append(i)
            if i < len(doc_copia.paragraphs):
                for run in doc_copia.paragraphs[i].runs:
                    run.font.color.rgb = RGBColor(0, 0, 255) # Resaltar en rojo el párrafo

    doc_copia.save("./documents/documento2_comparado.docx")

    if diferencias:
        return f"Se han encontrado diferencias. Las diferencias se han guardado en el documento: {output_path}."
    else:
        return "No se han encontrado diferencias entre los documentos."

# documento1_path = "documento1.docx"
# documento2_path = "documento2.docx"
# output_path = "documento1_diferencias.docx"
# output_path2 = "documento2_diferencias.docx"
# resaltar_diferencias(documento1_path, documento2_path, output_path)
# resaltar_diferencias_documento2(documento1_path, documento2_path, output_path2)


